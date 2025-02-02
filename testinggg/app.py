import requests
import os
import uuid
import logging
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import io
import contextlib

from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__, template_folder='.')

# Hyperbolic API Key
API_KEY = os.getenv("API_KEY") # Replace with actual API key
BASE_URL = "https://api.hyperbolic.xyz/v1/chat/completions"

app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['TEMP_FOLDER'] = os.path.join(os.path.dirname(__file__), 'temp')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMP_FOLDER'], exist_ok=True)

logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[
                        logging.StreamHandler(),
                        logging.FileHandler('app.log')
                    ])

def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, "rb") as pdf_file:
            reader = PdfReader(pdf_file)
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text()).strip()
    except Exception as e:
        logging.exception(f"Error extracting text from PDF: {e}")
        return f"Error extracting text from PDF: {e}"

def solve_coding_problem(question):
    try:
        prompt = f"""
        Write a Python solution for this problem: {question}.
        If the problem requires user input, assume random values instead.
        Provide only the code, no explanations or comments.
        skip words like assigment and dont make it a question.
        """
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {API_KEY}"
        }
        data = {
            "messages": [{"role": "user", "content": prompt}],
            "model": "meta-llama/Llama-3.3-70B-Instruct",
            "max_tokens": 512
        }
        response = requests.post(BASE_URL, headers=headers, json=data)
        response_json = response.json()
        solution = response_json.get("choices", [{}])[0].get("message", {}).get("content", "Error: No solution found.").strip()
        return solution.replace("```python", "").replace("```", "").strip()
    except Exception as e:
        logging.error(f"Error solving coding problem: {e}")
        return f"Error solving coding problem: {e}"

def execute_code(code):
    try:
        output = io.StringIO()
        with contextlib.redirect_stdout(output), contextlib.redirect_stderr(output):
            exec(code, {}, {})  
        return output.getvalue().strip()
    except Exception as e:
        logging.exception(f"Error executing code: {e}")
        return f"Error executing code: {e}"

def create_screenshot(output_text):
    try:
        output_text = output_text.strip() or "No output."
        
        font_path = "C:\\Windows\\Fonts\\consola.ttf"  # Windows example, adjust for other systems
        font = ImageFont.truetype(font_path, 14)

        lines = output_text.splitlines()
        width, height = 800, max(len(lines) * 16 + 20, 100)
        image = Image.new('RGB', (width, height), color='black')
        draw = ImageDraw.Draw(image)
        y_offset = 10
        
        for line in lines:
            draw.text((10, y_offset), line, fill='white', font=font)
            y_offset += 16
        
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        logging.exception(f"Error creating screenshot: {e}")
        return b""


def set_margins(doc, top=1, bottom=1, left=1, right=1):
    """Set document margins (in inches)."""
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def generate_word_doc(name, reg_number, questions, solutions, screenshots, output_path):
    try:
        doc = Document()

        # Set tighter margins to fit more on the page
        set_margins(doc, top=0.7, bottom=0.7, left=0.8, right=0.8)

        # Formatting styles
        heading_style = doc.styles['Heading 1']
        heading_style.font.size = Pt(16)

        subheading_style = doc.styles['Heading 2']
        subheading_style.font.size = Pt(14)

        code_style = doc.styles.add_style('CodeStyle', 1)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)

        # Add Name and Register Number
        doc.add_heading(f"Name: {name}", level=1)
        doc.add_heading(f"Register Number: {reg_number}", level=1)
        doc.add_paragraph("\n")  # Small spacer

        for i, (question, solution, screenshot) in enumerate(zip(questions, solutions, screenshots)):
            doc.add_heading(f"QUESTION {i + 1}", level=2)
            doc.add_paragraph(question)

            doc.add_heading("Code Solution", level=3)
            doc.add_paragraph(solution, style='CodeStyle')

            doc.add_heading("FINAL Output", level=3)

            # Shrink screenshot size for the first question to fit on the first page
            if screenshot:
                if i == 0:
                    doc.add_picture(io.BytesIO(screenshot), width=Inches(4.5))  # Smaller for first page
                else:
                    doc.add_picture(io.BytesIO(screenshot), width=Inches(6))  # Normal size for other pages

            # Add a page break after the first question to ensure clean separation
            if i == 0:
                doc.add_page_break()

        # Save the document
        doc.save(output_path)
        return output_path

    except Exception as e:
        logging.exception(f"Error generating Word document: {e}")
        return f"Error generating Word document: {e}"

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/solve', methods=['POST'])
def solve():
    try:
        name = request.form.get('name')
        reg_number = request.form.get('regNo')  # Fixed here
        file = request.files.get('file')

        if not file or not name or not reg_number:
            return jsonify({"error": "Name, Register Number, and file are required."}), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        pdf_text = extract_text_from_pdf(file_path)
        questions = pdf_text.split('\n')
        solutions = [solve_coding_problem(q) for q in questions]
        final_outputs = [execute_code(sol) for sol in solutions]
        screenshots = [create_screenshot(output) for output in final_outputs]

        output_file = os.path.join(app.config['TEMP_FOLDER'], f"solutions_{uuid.uuid4().hex}.docx")
        generate_word_doc(name, reg_number, questions, solutions, screenshots, output_file)

        return send_file(output_file, as_attachment=True)
    except Exception as e:
        logging.exception(f"A top-level error occurred: {e}")
        return jsonify({"error": "An unexpected error occurred."}), 500


